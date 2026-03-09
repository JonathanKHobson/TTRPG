from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/jonathanhobson/Downloads/Fantasy-War-Sim")
SOURCE = ROOT / "assets" / "documentation" / "War_Encounter_Rules_2026_v4.md"
OUTPUT_DOCX = ROOT / "output" / "doc" / "war-encounter-rules-2026-review-report.docx"
OUTPUT_MD = ROOT / "tmp" / "docs" / "war-encounter-rules-2026-review-report.md"

REPORT_TITLE = "War Encounter Rules 2026 Review Report"
REPORT_SUBTITLE = "Design critique and recommendations for the v4 markdown ruleset"
REPORT_DATE = "March 8, 2026"

REVIEW_CONTEXT = [
    ("Source file", "assets/documentation/War_Encounter_Rules_2026_v4.md"),
    ("Rules version", "Version date in document: 2026-02-24"),
    ("Primary mode", "Critique/Post-Playtest Diagnosis"),
    (
        "Secondary lenses",
        "Mechanics and Core Loop Design; Balance and Decision Analysis; Onboarding/Pacing/Readability",
    ),
    ("Audience", "Mixed design, product, and implementation collaborators"),
    ("Review stance", "Preserve the fast, GM-flexible philosophy; tighten only where ambiguity harms play"),
]

EXEC_SUMMARY = [
    "The ruleset has a strong high-level promise. The two-layer structure aims at a real D&D war problem: resolve the big war quickly while keeping PCs central to the story.",
    "The strongest ideas are the strategic-to-cinematic handoff, the emphasis on visible war consequences, and the attempt to keep large battles in a Risk-like resolution model instead of tactical miniatures accounting.",
    "The document is not yet stable enough for broad table use. Duplicate sections, drafting artifacts, terminology drift, off-page dependencies, and contradictory wording make the rules hard to trust on first read.",
    "The biggest design risk is stack overload. Army Size, STR, matchup effects, doctrines, meters, weather, terrain, help actions, and repeated attacks all compete inside the same battle math, which raises GM burden and flattens some intended distinctions.",
    "Recommended direction: stabilize the text first, standardize the cinematic impact ladder, reduce battle-math compression, and publish the system in explicit play tiers rather than treating every subsystem as equally active by default.",
]

PROMISE_ROWS = [
    (
        "Fantasy and promise",
        "Resolve large-scale war quickly without turning PCs into spectators.",
        "Strong target. The concept section consistently aims at fast war, visible stakes, and hero agency.",
    ),
    (
        "Core loop",
        "Move -> declare -> optional cinematic -> battle -> aftermath.",
        "Promising, but overloaded once optional systems and repeated attacks are layered in.",
    ),
    (
        "Meaningful play",
        "PC actions should clearly change strategic outcomes.",
        "Partially successful. The document promises impact, but the conversion from scene success to battle effect is still too discretionary.",
    ),
    (
        "Readability",
        "A GM should be able to run a war stage with confidence.",
        "Currently weak. The system needs a stabilization pass before its best ideas can land cleanly at the table.",
    ),
]

WHAT_WORKS_WELL = [
    "The overview chapters clearly identify the player-facing fantasy: fast war, real stakes, and personal hero relevance.",
    "The two-layer split is the right structural choice for a D&D-adjacent war system. It lets the war move without forcing every clash into tactical combat.",
    "The strategic loop is short enough to be teachable in concept. Move, declare, optionally zoom in, then resolve battles is the right backbone.",
    "Army classes have strong thematic silhouettes. Even before tuning, players can infer what Archers, Cavalry, Pikes, Siege, Spies, and Mages are supposed to do.",
    "Doctrines and command meters create real campaign texture. They point toward a compelling version of the game where logistics, morale, and politics matter without a full grand strategy sim.",
    "The weather and terrain material is evocative and campaign-friendly. Once trimmed and modularized, it can become a strong amplifier for scenario identity.",
]

FINDINGS = [
    {
        "severity": "Critical",
        "title": "Editorial instability makes the rules hard to trust",
        "where": "Chapters 6-10 and multiple cross-references",
        "diagnosis": (
            "The rules file still contains duplicated sections, visible drafting artifacts, terminology drift, and at least one live dependency on an external Google Doc. Examples include duplicated 'Invading a Location', 'Supply Meter', and 'Influence Meter' sections; drafting notes such as 'Perfect. We'll mirror...' and 'Absolutely. Here's a cleaned + expanded version...'; and word drift such as 'Moral Meter' instead of 'Morale Meter'."
        ),
        "rewards": (
            "The current text rewards improvisation over reliable procedure. A GM is pushed to infer the intended rule rather than follow a stable source of truth."
        ),
        "recommendation": (
            "Freeze a master rulebook, remove all duplicated text and drafting artifacts, eliminate off-page dependencies, and run a terminology/style pass before further balance expansion. This is the highest-value change because every later tuning pass depends on readers trusting the document."
        ),
        "tradeoff": (
            "This pauses new content for one editorial cycle, but it creates the foundation that makes later balance work worth doing."
        ),
    },
    {
        "severity": "Critical",
        "title": "Core battle math compresses differentiation and makes the cap do too much work",
        "where": "Chapter 4 and Chapter 6 battle pool rules",
        "diagnosis": (
            "Army Size and STR both add full dice, then attacker and defender caps flatten the result. The worked example reaches 15 vs 10 dice before modifiers, then collapses to 10 vs 8 after caps. That means large chunks of army identity, terrain, and doctrine value can become invisible because many inputs compete for the same final two dice."
        ),
        "rewards": (
            "The system currently rewards stacking raw STR and only then looking for a few decisive post-cap penalties or bonuses. It is less about nuanced army identity than about hitting the cap efficiently."
        ),
        "recommendation": (
            "Separate base force from situational edge. Let Army Size drive most of the pool, reduce STR to smaller quality bands or die-step shifts, and reserve only a narrow post-cap bonus window for terrain, matchup, doctrine, and cinematic impact. This keeps the roll fast while making modifiers legible again."
        ),
        "tradeoff": (
            "This requires reworking examples and some class assumptions, but it will make the system more readable and less prone to hidden saturation."
        ),
    },
    {
        "severity": "Critical",
        "title": "Optional systems are not truly modular and overload the GM's decision surface",
        "where": "Chapters 7-10 plus optional mechanics inside Chapter 6",
        "diagnosis": (
            "Meters, doctrines, help actions, medals, financial rules, weather, terrain, and campaign economics all plug into the same battle resolution space, but the document does not clearly define which of these are expected in a default game. The result is a rulebook that reads as if everything might matter at once."
        ),
        "rewards": (
            "The current structure rewards veteran GMs who can prune subsystems on the fly. It punishes first-time readers with a high cognitive load before the core loop is even stable."
        ),
        "recommendation": (
            "Publish explicit play tiers: Core War, Command War, and Full Campaign. Each subsystem should declare which mode uses it and which phase it modifies. This turns optional material into real modules instead of ambient overhead."
        ),
        "tradeoff": (
            "The document becomes slightly more structured and less freeform, but it becomes far easier to teach and adopt."
        ),
    },
    {
        "severity": "Major",
        "title": "The cinematic-to-strategic handoff is thematically strong but mechanically under-specified",
        "where": "Quick Start, Chapter 1, Chapter 2, and battle pool construction",
        "diagnosis": (
            "The rules correctly frame cinematic sequences as the place where PCs matter, but the conversion language is too loose. Outcomes can add or remove dice, change die sizes, alter conditions, or occasionally decide a battle entirely, yet there is no standard effect ladder, declared-stakes template, or cap on how many strategic changes one scene can generate."
        ),
        "rewards": (
            "The current system rewards GM fiat and player persuasion more than predictable agency. Players cannot estimate the strategic value of choosing one intervention over another."
        ),
        "recommendation": (
            "Predeclare each cinematic objective and map it to a named reward tier before the scene starts: minor, major, or decisive. Keep decisive outcomes rare and only for predeclared climax objectives."
        ),
        "tradeoff": (
            "This narrows improvisational range slightly, but it improves player autonomy and consequence clarity."
        ),
    },
    {
        "severity": "Major",
        "title": "Battle cadence currently favors deterministic grinding and rout spirals",
        "where": "Chapter 6 battle aftermath, retreat, and repeated attack rules",
        "diagnosis": (
            "Attackers can press repeated attacks in the same Battle Phase, defenders auto-lose when their pool hits zero, and retreat relies on a flat 1-in-6 roll after attacks. Combined with a five-die minimum to initiate and ties favoring defenders, this creates a swingy but still grind-forward cadence where the first strong engagement often dictates the rest of the phase."
        ),
        "rewards": (
            "The rules currently reward concentration of attack power and repeated pressing over repositioning, delaying, or elastic defense."
        ),
        "recommendation": (
            "After the first loss, defenders should explicitly choose hold, withdraw, or desperate stand. Additional presses should cost exhaustion, CP, or doctrine support. Retreat odds should depend on mobility and terrain rather than a flat 1d6 for every army."
        ),
        "tradeoff": (
            "Some clashes will take one extra decision step, but wars will feel less pre-solved after the first exchange."
        ),
    },
    {
        "severity": "Major",
        "title": "Army and doctrine abilities use inconsistent templates, hiding counterplay and stack rules",
        "where": "Chapter 5, Chapter 6 optional mechanics, and Chapter 9 doctrines",
        "diagnosis": (
            "High-leverage abilities vary widely in timing, range, visibility, frequency, and counter language. Assassins remove the highest enemy die or directly bleed size while unseen. Spies are effectively always hidden except to nearby spies. Heavy-Protection grants near-universal resistance. Warforged change their entire die type by matchup. These all create meaningful identity, but the template is inconsistent enough that balance and dispute resolution become difficult."
        ),
        "rewards": (
            "The current text rewards exploiting loosely specified edge cases and asking the GM for favorable interpretations."
        ),
        "recommendation": (
            "Standardize every class and doctrine entry around the same fields: phase, trigger, range, frequency, effect type, counterplay, and stack rule. If an ability modifies pool size, die size, or outcome timing, say so directly."
        ),
        "tradeoff": (
            "Entries become slightly longer, but table rulings become faster and more defensible."
        ),
    },
    {
        "severity": "Major",
        "title": "The economy and governance layer is appealing but still incomplete",
        "where": "Chapter 6 financial rules and Chapter 7 command systems",
        "diagnosis": (
            "CP is presented as a scarce, high-value war resource, but its gain cadence and default sinks remain partly descriptive rather than procedural. Medals, improved training, the Republic of Fairgard ledger, and some construction or policy references assume adjacent systems that are not fully present in this document."
        ),
        "rewards": (
            "The current structure rewards either full GM invention or complete subsystem avoidance. Neither is a good default for a first release."
        ),
        "recommendation": (
            "Either add a one-page economy loop with default numbers and turn timing, or move incomplete economic hooks into a campaign appendix until they are fully implemented."
        ),
        "tradeoff": (
            "The first stable release may be narrower, but it will be more coherent and easier to tune."
        ),
    },
    {
        "severity": "Moderate",
        "title": "Chapter 5 mixes mechanical archetypes, species rules, and named factions in a way that slows onboarding",
        "where": "Chapter 5 army classifications and faction sections",
        "diagnosis": (
            "The army catalog combines generic battlefield roles, stealth specialists, supernatural species, and setting-specific rosters in one continuous lookup block. That makes it harder to distinguish reusable core mechanics from setting flavor or campaign-specific exceptions."
        ),
        "rewards": (
            "The current organization rewards designer familiarity more than reader comprehension."
        ),
        "recommendation": (
            "Split the roster into mechanical archetypes, species or special templates, and faction overlays. Keep a single stat-entry template across all three."
        ),
        "tradeoff": (
            "This adds one more table of contents layer, but it makes the rules far easier to learn and expand."
        ),
    },
    {
        "severity": "Moderate",
        "title": "Weather and terrain are evocative, but the chapter is too broad for the current core loop",
        "where": "Chapter 10 and class preferred/undesired environment hooks",
        "diagnosis": (
            "The environmental material has good flavor and useful scenario hooks, but it introduces a large lookup surface on top of already dense battle math. Because so many armies also have preferred or undesired weather and terrain, every battle risks becoming a multi-table rules audit."
        ),
        "rewards": (
            "The current version rewards using only the most obvious environment effects while quietly ignoring the rest."
        ),
        "recommendation": (
            "Publish a starter environment kit of six to eight common conditions, then keep the long weather and terrain list as an advanced appendix. Use simple tags for modifier handoff."
        ),
        "tradeoff": (
            "The core game loses some encyclopedic breadth, but gains much faster battle setup."
        ),
    },
    {
        "severity": "Minor",
        "title": "The quick reference cannot become reliable until terminology and chapter text are stabilized",
        "where": "Appendix and cross-document references",
        "diagnosis": (
            "The appendix is useful in intent, but a quick reference only works when the main rule language is already normalized. Right now the book still contains wording drift, duplicate sections, and chapter overlap that prevent the appendix from serving as a trustworthy compression layer."
        ),
        "rewards": (
            "The current appendix rewards experienced readers who already know what the system probably means."
        ),
        "recommendation": (
            "Rebuild the appendix after the chapter text is stabilized. Treat it as a final product, not an early drafting scaffold."
        ),
        "tradeoff": (
            "This delays final reference polish, but avoids reinforcing unstable rules wording."
        ),
    },
]

CHANGE_SEQUENCE = [
    "Run an editorial stabilization pass first. Remove duplicated sections, drafting notes, broken references, and terminology drift before touching balance numbers.",
    "Redesign the battle math around fewer meaningful knobs: clear base force, clear quality modifier, and a small post-cap window for situational effects.",
    "Add a cinematic impact ladder with predeclared stakes so players can predict how hero scenes matter strategically.",
    "Publish the rules in three modes - Core War, Command War, and Full Campaign - and mark every subsystem with the modes that use it.",
    "Standardize all army and doctrine entries around the same timing and counterplay template.",
    "Either complete the economy layer with default cadence and numbers, or demote unfinished pieces to campaign hooks until they are ready.",
]

REWRITE_SAMPLES = [
    {
        "title": "Suggested replacement text: Cinematic impact ladder",
        "intent": "Replace ad hoc scene-to-strategy conversion with a bounded ladder.",
        "lines": [
            "When the GM declares a Cinematic Sequence, name the objective and its reward tier before play begins.",
            "Minor objective: on success, apply one of the following to the next related battle only: +1 die, -1 die, Advantage 1, or cancel one minor terrain penalty.",
            "Major objective: on success, apply one of the following to the next related battle only: +2 dice, -2 dice, upgrade or downgrade one die step on up to two dice, or disable one named doctrine or fortification benefit.",
            "Decisive objective: only for predeclared climax scenes. On success, force a retreat, cancel the battle entirely, or remove a named commander or special asset from the battle. A decisive reward must replace normal dice rewards, not stack with them.",
            "A single Cinematic Sequence can grant only one reward tier unless the GM declares that both sides entered the scene with separate objectives.",
        ],
    },
    {
        "title": "Suggested replacement text: Battle pool construction",
        "intent": "Reduce cap compression while keeping the roll fast.",
        "lines": [
            "Build each army's battle pool in this order.",
            "1. Base force: gain 1 die per 50 Army Size, rounded down.",
            "2. Quality band: compare the army's STR to the war's baseline and apply one result - Weak: -1 die, Standard: no change, Veteran: +1 die, Elite: +2 dice or upgrade one die step on two dice.",
            "3. Situational edge: apply no more than two sources from terrain, weather, matchup, doctrine, help, or cinematic results.",
            "4. Final cap: Attacker 8 dice, Defender 7 dice. If an army would gain more dice after reaching its cap, convert each two excess dice into Advantage 1, to a maximum of Advantage 2.",
            "This keeps large armies dangerous, preserves situational edges, and prevents late modifiers from disappearing into the cap.",
        ],
    },
    {
        "title": "Suggested replacement text: Trade doctrine trigger template",
        "intent": "Normalize doctrine timing and avoid repeated ambiguous phrasing.",
        "lines": [
            "Trigger timing: resolve doctrine upkeep, earnings, and meter shifts at the start of the Move Phase.",
            "Failure check: if a doctrine lists a failure threshold and that meter is at or below the threshold at this moment, the doctrine stays active but its positive effect is suppressed for this turn.",
            "Penalty rule: apply the doctrine's listed penalties once for the turn. Do not reapply them again when the doctrine is deactivated later in the same turn.",
            "Naming rule: always use the same meter names - Morale, Supply, and Influence - in both doctrine text and chapter text.",
        ],
    },
    {
        "title": "Suggested replacement text: Rules mode gate",
        "intent": "Turn optional material into true modules.",
        "lines": [
            "Choose one war mode before the campaign begins.",
            "Core War: use Quick Start, Chapters 1-6, and the appendix only.",
            "Command War: add Chapter 7 command meters and any doctrines explicitly marked Command War.",
            "Full Campaign: add Chapter 10 environment rules, financial rules, medals, and advanced doctrines.",
            "Every optional rule must list which war mode uses it and which phase it modifies.",
        ],
    },
]

PLAYTESTS = [
    {
        "name": "Core loop teach test",
        "question": "Can a new GM run one full war stage using only Quick Start plus the stabilized Chapter 6 procedure?",
        "build": "Use the core loop only. No meters, no doctrines, no weather table beyond one obvious condition.",
        "scenario": "Two armies on a simple map with one defended location and one possible cinematic intervention.",
        "observe": [
            "How many rules lookups happen outside Quick Start and Chapter 6",
            "Whether the GM can explain why the battle result happened",
            "Whether players can predict what their choices will change",
        ],
        "success": "The GM resolves the stage with two or fewer outside lookups and players can restate the battle logic in plain language.",
        "failure": "The GM needs ad hoc rulings for core timing, pool math, or cinematic conversion.",
    },
    {
        "name": "Cinematic agency test",
        "question": "Do players feel they are choosing where to matter, rather than just being handed a spotlight scene?",
        "build": "Run two possible cinematic objectives with predeclared reward tiers tied to different fronts.",
        "scenario": "Players must choose whether to sabotage a siege engine or rescue a commander before the battle phase.",
        "observe": [
            "Whether players understand the tradeoff before choosing",
            "Whether the chosen scene produces a clear strategic change",
            "Whether the unchosen front still feels like a meaningful cost",
        ],
        "success": "Players report that the choice felt legible and consequential before the dice were rolled.",
        "failure": "Players describe the scene as cool but arbitrary, or cannot explain why one objective mattered more than the other.",
    },
    {
        "name": "Cap compression test",
        "question": "Do different armies still feel different after pool caps are applied?",
        "build": "Run at least twelve sample battles across low, mid, and high force levels with and without terrain or doctrine modifiers.",
        "scenario": "Include cavalry vs pikes, siege vs fortification, warforged vs anti-magic opposition, and one stealth-heavy matchup.",
        "observe": [
            "How often different armies collapse to the same final pool",
            "Whether terrain and doctrine choices still move outcomes materially",
            "Which modifiers routinely become invisible after capping",
        ],
        "success": "Distinct army profiles still produce distinct final states after the cap, and situational choices move the outcome often enough to matter.",
        "failure": "Most battles saturate into the same capped pools and only a few extreme modifiers remain relevant.",
    },
    {
        "name": "Module burden test",
        "question": "Do advanced modules add texture without breaking pacing?",
        "build": "Run the same scenario once in Core War mode and once in Full Campaign mode.",
        "scenario": "A three-front war turn with one doctrine, one weather condition, and one command meter swing.",
        "observe": [
            "Resolution time",
            "Number of rule lookups and disputes",
            "Whether added systems create new decisions or just more bookkeeping",
        ],
        "success": "Advanced play produces new decisions without doubling resolution friction.",
        "failure": "The advanced version mostly adds arithmetic, forgotten triggers, and delayed payoff.",
    },
    {
        "name": "Counterplay integrity test",
        "question": "Do signature units and doctrines have readable answers rather than hard invalidation or no answer at all?",
        "build": "Stage focused matchups for Vampires, Assassins, Warforged, Siege, and Heavy-Protection armies against prepared counters.",
        "scenario": "Use one prepared counter and one unprepared control case for each featured unit or doctrine.",
        "observe": [
            "Whether the counter changes the odds in a meaningful but not automatic way",
            "Whether the defending player understands how to access the counterplay",
            "Whether the signature unit still feels special after being answered",
        ],
        "success": "Counters create pressure and adaptation without deleting the fantasy of the signature unit.",
        "failure": "A counter fully blanks the unit, or the unit remains dominant despite the supposed answer.",
    },
]

APPENDIX_ROWS = [
    (
        "Quick Start",
        "Clear promise and usable war-loop summary.",
        "Needs tighter link to the later chapter language so it remains a true entry point.",
        "Major",
    ),
    (
        "Chapter 1 - Overview",
        "Best articulation of fantasy, player role, and two-layer logic.",
        "Needs a firmer rule for when cinematic play changes strategy.",
        "Major",
    ),
    (
        "Chapter 2 - Cinematic Sequences",
        "Strong PC spotlight intent.",
        "Needs a standard reward ladder and clearer bounds on narrative overrides.",
        "Major",
    ),
    (
        "Chapter 3 - Size Classifications",
        "Useful framing for cinematic NPC scale.",
        "Needs tighter linkage to the strategic layer and fewer one-off exceptions.",
        "Moderate",
    ),
    (
        "Chapter 4 - Strategic Wargame Sequences",
        "Good conceptual frame and readable high-level procedure.",
        "Battle math explanation overpromises speed while hiding saturation risk.",
        "Critical",
    ),
    (
        "Chapter 5 - Army Classifications and Factions",
        "Strong thematic silhouettes and faction flavor.",
        "Lookup burden, template inconsistency, and taxonomy sprawl slow onboarding.",
        "Major",
    ),
    (
        "Chapter 6 - Core Procedures",
        "Contains the real backbone of the system.",
        "Repeated attack cadence, retreat rules, and pool construction need revision.",
        "Critical",
    ),
    (
        "Chapter 7 - Strategic Command Meters",
        "Adds compelling campaign texture.",
        "Contains duplicate headings and pushes too much load into the default loop.",
        "Critical",
    ),
    (
        "Chapter 8 - Optional Mechanics",
        "Contains useful hooks such as help actions and environment toggles.",
        "Needs mode gating so optional rules stop feeling mandatory.",
        "Major",
    ),
    (
        "Chapter 9 - Doctrines",
        "Rich strategic identity and strong flavor scaffolding.",
        "Power inflation, ambiguous timing, and inconsistent templates undermine tuning.",
        "Major",
    ),
    (
        "Chapter 10 - Weather and Terrain",
        "Evocative battlefield flavor and scenario texture.",
        "Too broad for core use and still contains drafting artifacts.",
        "Moderate",
    ),
    (
        "Appendix",
        "Correct instinct: quick references matter for this kind of game.",
        "Should be rebuilt only after the main text is stabilized.",
        "Minor",
    ),
]

ASSUMPTIONS = [
    "This review evaluates the markdown rules file as a document for table use, not the underlying HTML or app implementation.",
    "The intended format is a GM-facing TTRPG or wargame hybrid, not a competitive tournament ruleset.",
    "Recommendations prioritize clarity, pacing, and agency over simulation detail.",
    "The immediate goal is a stable v1 rulebook, not exhaustive setting coverage.",
]


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_padding(cell, top=48, bottom=48, left=54, right=54):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure_styles(document):
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10)

    for style_name, size in [("Heading 1", 13), ("Heading 2", 11.5), ("Heading 3", 10.5)]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(size)
        style.font.bold = True


def add_title_block(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=2, line=1.0)
    run = title.add_run(REPORT_TITLE)
    run.bold = True
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=1, line=1.0)
    run = subtitle.add_run(REPORT_SUBTITLE)
    run.italic = True
    run.font.size = Pt(10)

    date_line = document.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(date_line, after=8, line=1.0)
    run = date_line.add_run(f"Prepared {REPORT_DATE}")
    run.font.size = Pt(9)


def add_heading(document, text, level):
    paragraph = document.add_heading(text, level=level)
    if level == 1:
        set_paragraph_spacing(paragraph, before=6, after=3, line=1.0)
    elif level == 2:
        set_paragraph_spacing(paragraph, before=5, after=2, line=1.0)
    else:
        set_paragraph_spacing(paragraph, before=3, after=1, line=1.0)
    return paragraph


def add_paragraph(document, text, *, italic=False):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=3, line=1.1)
    run = paragraph.add_run(text)
    run.italic = italic
    run.font.size = Pt(10)
    return paragraph


def add_labeled_paragraph(document, label, text):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=2, line=1.1)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    body_run = paragraph.add_run(text)
    body_run.font.size = Pt(10)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, after=1, line=1.08)
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        run = paragraph.add_run(item)
        run.font.size = Pt(9.8)


def add_numbered_steps(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        set_paragraph_spacing(paragraph, after=2, line=1.08)
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.first_line_indent = Inches(-0.14)
        run = paragraph.add_run(item)
        run.font.size = Pt(9.8)


def add_simple_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].width = Inches(widths[idx])
        set_cell_padding(hdr[idx], top=56, bottom=56)
        shade_cell(hdr[idx], "D9E8F5")
        para = hdr[idx].paragraphs[0]
        set_paragraph_spacing(para, after=0, line=1.0)
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)

    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].width = Inches(widths[idx])
            set_cell_padding(row[idx])
            para = row[idx].paragraphs[0]
            set_paragraph_spacing(para, after=0, line=1.0)
            run = para.add_run(value)
            run.font.size = Pt(9.1)
    return table


def add_rewrite_block(document, sample):
    add_heading(document, sample["title"], 2)
    add_labeled_paragraph(document, "Intent", sample["intent"])
    block = document.add_paragraph()
    set_paragraph_spacing(block, after=4, line=1.0)
    block.paragraph_format.left_indent = Inches(0.18)
    block.paragraph_format.right_indent = Inches(0.08)
    for idx, line in enumerate(sample["lines"]):
        run = block.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        if idx < len(sample["lines"]) - 1:
            run.add_break()


def build_markdown():
    lines = [
        f"# {REPORT_TITLE}",
        "",
        REPORT_SUBTITLE,
        "",
        f"Prepared {REPORT_DATE}",
        "",
        "## Executive summary",
        "",
    ]
    lines.extend([f"- {item}" for item in EXEC_SUMMARY])
    lines.extend(["", "## Review context", ""])
    for key, value in REVIEW_CONTEXT:
        lines.append(f"- **{key}:** {value}")

    lines.extend(["", "## System promise and intended player experience", ""])
    lines.append("| Lens | Intended outcome | Assessment |")
    lines.append("| --- | --- | --- |")
    for row in PROMISE_ROWS:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} |")

    lines.extend(["", "## What is working well", ""])
    lines.extend([f"- {item}" for item in WHAT_WORKS_WELL])

    lines.extend(["", "## Priority findings overview", ""])
    lines.append("| Severity | Finding | Area | Main move |")
    lines.append("| --- | --- | --- | --- |")
    for finding in FINDINGS:
        lines.append(
            f"| {finding['severity']} | {finding['title']} | {finding['where']} | {finding['recommendation']} |"
        )

    lines.extend(["", "## Priority findings", ""])
    for finding in FINDINGS:
        lines.extend(
            [
                f"### {finding['severity']} - {finding['title']}",
                "",
                f"- **Where it appears:** {finding['where']}",
                f"- **Diagnosis:** {finding['diagnosis']}",
                f"- **What the system rewards today:** {finding['rewards']}",
                f"- **Recommendation:** {finding['recommendation']}",
                f"- **Tradeoff:** {finding['tradeoff']}",
                "",
            ]
        )

    lines.extend(["## Recommendations and targeted fix proposals", ""])
    lines.extend([f"1. {item}" for item in CHANGE_SEQUENCE])

    lines.extend(["", "## Suggested rewrite text for the highest-impact rules only", ""])
    for sample in REWRITE_SAMPLES:
        lines.extend([f"### {sample['title']}", "", f"**Intent:** {sample['intent']}", "", "```text"])
        lines.extend(sample["lines"])
        lines.extend(["```", ""])

    lines.extend(["## Playtest and validation plan", ""])
    for test in PLAYTESTS:
        lines.extend(
            [
                f"### {test['name']}",
                "",
                f"- **Question:** {test['question']}",
                f"- **Build under test:** {test['build']}",
                f"- **Scenario:** {test['scenario']}",
                "- **What to observe:**",
            ]
        )
        lines.extend([f"  - {item}" for item in test["observe"]])
        lines.extend(
            [
                f"- **Success signal:** {test['success']}",
                f"- **Failure signal:** {test['failure']}",
                "",
            ]
        )

    lines.extend(["## Appendix: issue map by chapter", ""])
    lines.append("| Area | Strength | Main risk | Priority |")
    lines.append("| --- | --- | --- | --- |")
    for row in APPENDIX_ROWS:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} |")

    lines.extend(["", "## Assumptions", ""])
    lines.extend([f"- {item}" for item in ASSUMPTIONS])
    return "\n".join(lines) + "\n"


def build_docx():
    document = Document()
    configure_styles(document)
    add_title_block(document)

    add_heading(document, "Executive Summary", 1)
    add_bullets(document, EXEC_SUMMARY)

    add_heading(document, "Review Context", 1)
    add_simple_table(document, ["Field", "Value"], REVIEW_CONTEXT, [1.7, 4.95])

    add_heading(document, "System Promise and Intended Player Experience", 1)
    add_simple_table(
        document,
        ["Lens", "Intended outcome", "Assessment"],
        PROMISE_ROWS,
        [1.45, 2.15, 3.05],
    )

    add_heading(document, "What Is Working Well", 1)
    add_bullets(document, WHAT_WORKS_WELL)

    add_heading(document, "Priority Findings Overview", 1)
    overview_rows = [
        (finding["severity"], finding["title"], finding["where"], finding["recommendation"])
        for finding in FINDINGS
    ]
    add_simple_table(
        document,
        ["Severity", "Finding", "Area", "Main recommendation"],
        overview_rows,
        [0.95, 2.25, 1.6, 2.55],
    )

    add_heading(document, "Priority Findings", 1)
    for finding in FINDINGS:
        add_heading(document, f"{finding['severity']} - {finding['title']}", 2)
        add_labeled_paragraph(document, "Where it appears", finding["where"])
        add_labeled_paragraph(document, "Diagnosis", finding["diagnosis"])
        add_labeled_paragraph(document, "What the system rewards today", finding["rewards"])
        add_labeled_paragraph(document, "Recommendation", finding["recommendation"])
        add_labeled_paragraph(document, "Tradeoff", finding["tradeoff"])

    add_heading(document, "Recommendations and Targeted Fix Proposals", 1)
    add_numbered_steps(document, CHANGE_SEQUENCE)

    add_heading(document, "Suggested Rewrite Text for the Highest-Impact Rules Only", 1)
    for sample in REWRITE_SAMPLES:
        add_rewrite_block(document, sample)

    add_heading(document, "Playtest and Validation Plan", 1)
    for test in PLAYTESTS:
        add_heading(document, test["name"], 2)
        add_labeled_paragraph(document, "Question", test["question"])
        add_labeled_paragraph(document, "Build under test", test["build"])
        add_labeled_paragraph(document, "Scenario", test["scenario"])
        add_labeled_paragraph(document, "What to observe", "")
        add_bullets(document, test["observe"])
        add_labeled_paragraph(document, "Success signal", test["success"])
        add_labeled_paragraph(document, "Failure signal", test["failure"])

    add_heading(document, "Appendix: Issue Map by Chapter", 1)
    add_simple_table(
        document,
        ["Area", "Strength", "Main risk", "Priority"],
        APPENDIX_ROWS,
        [1.35, 1.95, 2.95, 0.7],
    )

    add_heading(document, "Assumptions", 1)
    add_bullets(document, ASSUMPTIONS)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)


def validate_docx():
    document = Document(OUTPUT_DOCX)
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    required_strings = [
        REPORT_TITLE,
        "Executive Summary",
        "Priority Findings",
        "Playtest and Validation Plan",
        "Appendix: Issue Map by Chapter",
    ]
    missing = [item for item in required_strings if item not in text]
    if missing:
        raise RuntimeError(f"Missing expected document content: {missing}")
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "missing": missing,
    }


def main():
    markdown = build_markdown()
    OUTPUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_MD.write_text(markdown)
    build_docx()
    result = validate_docx()
    print(OUTPUT_DOCX)
    print(OUTPUT_MD)
    print(f"paragraphs={result['paragraphs']}")
    print(f"tables={result['tables']}")


if __name__ == "__main__":
    main()
