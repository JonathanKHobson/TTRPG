from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/jonathanhobson/Downloads/Fantasy-War-Sim")
OUTPUT = ROOT / "output" / "doc" / "fantasy-war-sim-one-page-summary.docx"


WHAT_IT_IS = (
    "Fantasy War Sim is a local-first fantasy battle simulator with transparent dice "
    "math, War Report staging and auto-battle flows, saved armies and factions, battle "
    "history, and browser-only persistence. The shipped app is an offline-ready static "
    "web app with no backend or build step required for normal use."
)

WHO_ITS_FOR = "Not found in repo."

FEATURES = [
    "Run local fantasy battles between attacker and defender armies.",
    "Load saved armies or create new ones directly in the battle flow.",
    "Save and manage army and faction data inside the app.",
    "Inspect transparent dice math through pool summaries and breakdown views.",
    "Use War Report staging for round setup, replay, and auto-battle control.",
    "Review grouped battle history by session and turn, with export and clear actions.",
    "Persist settings, weather, armies, factions, and history in browser localStorage.",
]

HOW_IT_WORKS = [
    "Main shell: fantasy-war-sim.html contains the UI structure, event wiring, runtime state, and battle logic.",
    "Styling: assets/css/war-table-theme.css provides the base theme; assets/css/war-table-state-themes.css layers runtime weather, phenomena, and defeat visuals.",
    "Supporting surfaces: fantasy-war-sim-help-wiki.html and War_Encounter_Rules_2026_v4.html are separate static HTML references with their own stylesheet layering.",
    "Data flow: user input updates runtime state, battle resolution runs locally, then results and history are persisted to localStorage under warTableState and re-rendered.",
]

HOW_TO_RUN = [
    "Open fantasy-war-sim.html in a browser.",
    "Open fantasy-war-sim-help-wiki.html for quick rules and help.",
    "Run npx playwright test --project=chromium-local for the local regression suite.",
]


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_padding(cell, top=24, bottom=24):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "bottom": bottom, "left": 36, "right": 36}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_body_paragraph(document, text, *, bold=False):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=2, line=1.0)
    run = paragraph.add_run(text)
    run.font.size = Pt(9.5)
    run.bold = bold
    return paragraph


def add_heading(document, text):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=2, after=1, line=1.0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, after=0, line=1.0)
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        run = paragraph.add_run(item)
        run.font.size = Pt(9.3)


def build_doc():
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=2, line=1.0)
    title_run = title.add_run("Fantasy War Sim")
    title_run.bold = True
    title_run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=6, line=1.0)
    subtitle_run = subtitle.add_run("One-page app summary")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(9)

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(5.95)

    sections = [
        ("What it is", WHAT_IT_IS),
        ("Who it's for", WHO_ITS_FOR),
        ("What it does", FEATURES),
        ("How it works", HOW_IT_WORKS),
        ("How to run", HOW_TO_RUN),
    ]

    for heading, content in sections:
        row = table.add_row()
        left = row.cells[0]
        right = row.cells[1]
        set_cell_padding(left)
        set_cell_padding(right)

        left_para = left.paragraphs[0]
        set_paragraph_spacing(left_para, after=0, line=1.0)
        left_run = left_para.add_run(heading)
        left_run.bold = True
        left_run.font.size = Pt(10)

        if isinstance(content, str):
            para = right.paragraphs[0]
            set_paragraph_spacing(para, after=0, line=1.0)
            run = para.add_run(content)
            run.font.size = Pt(9.3)
        else:
            first_para = right.paragraphs[0]
            first_para._element.getparent().remove(first_para._element)
            first_para._p = first_para._element = None
            for item in content:
                bullet = right.add_paragraph(style="List Bullet")
                set_paragraph_spacing(bullet, after=0, line=1.0)
                bullet.paragraph_format.left_indent = Inches(0.16)
                bullet.paragraph_format.first_line_indent = Inches(-0.11)
                run = bullet.add_run(item)
                run.font.size = Pt(9.1)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
